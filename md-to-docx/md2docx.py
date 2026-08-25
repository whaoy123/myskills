#!/usr/bin/env python3
"""Template-aware Markdown to DOCX converter.

Supported focus: engineering/course/academic Markdown with headings, paragraphs,
lists, pipe tables, fenced code, images, rendered Mermaid diagrams, captions,
and numeric bibliography references.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PAGE_W = 14.66
BOOKMARK_ID = 1000
CITATION_RE = re.compile(r"\[((?:\d+|[,\s，、;；\-–—])+)\]")
IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")
CAPTION_RE = re.compile(r"^(表|图)\s*\d+[\-.]\d+\s*(.*)")


def style_exists(doc, style_name: str) -> bool:
    try:
        doc.styles[style_name]
        return True
    except KeyError:
        return False


def set_paragraph_style(paragraph, doc, style_name: str) -> None:
    if style_exists(doc, style_name):
        paragraph.style = doc.styles[style_name]


def load_document(template_path: Path | None):
    """Open a .docx/.dotx template, or return a blank document."""
    if template_path is None:
        return Document()
    suffix = template_path.suffix.lower()
    if suffix == ".docx":
        return Document(str(template_path))
    if suffix != ".dotx":
        raise ValueError("template must be .docx or .dotx")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        tmp_path = Path(handle.name)
    try:
        shutil.copy(template_path, tmp_path)
        with zipfile.ZipFile(tmp_path, "r") as src:
            files = {}
            for item in src.infolist():
                data = src.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.decode("utf-8").replace(
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                    ).encode("utf-8")
                files[item.filename] = data
        tmp_path.unlink(missing_ok=True)
        with zipfile.ZipFile(tmp_path, "w") as dst:
            for name, data in files.items():
                dst.writestr(name, data)
        return Document(str(tmp_path))
    finally:
        tmp_path.unlink(missing_ok=True)


def usable_page_width_cm(doc) -> float:
    if not doc.sections:
        return 14.66
    section = doc.sections[0]
    width = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    return round(width, 2) if width > 0 else 14.66


def set_run_font(run, cn_font="宋体", en_font="Times New Roman", size=Pt(12)):
    run.font.name = en_font
    run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), cn_font)


def add_run_with_font(paragraph, text, cn_font="宋体", en_font="Times New Roman", size=Pt(12)):
    run = paragraph.add_run(text)
    set_run_font(run, cn_font, en_font, size)
    return run


def ref_bookmark_name(number: str) -> str:
    return f"ref_{number}"


def next_bookmark_id() -> str:
    global BOOKMARK_ID
    BOOKMARK_ID += 1
    return str(BOOKMARK_ID)


def set_citation_run_style(run, size=Pt(9)):
    set_run_font(run, size=size)
    run.font.superscript = True


def add_bookmarked_reference_label(paragraph, number: str):
    bookmark_id = next_bookmark_id()
    open_run = paragraph.add_run("[")
    set_run_font(open_run)

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), ref_bookmark_name(number))
    paragraph.add_run()._element.append(start)

    num_run = paragraph.add_run(number)
    set_run_font(num_run)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph.add_run()._element.append(end)

    close_run = paragraph.add_run("] ")
    set_run_font(close_run)


def add_ref_field(paragraph, number: str, size=Pt(9)):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run = paragraph.add_run()
    set_citation_run_style(run, size)
    run._element.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" REF {ref_bookmark_name(number)} \\h \\* CHARFORMAT "
    run = paragraph.add_run()
    set_citation_run_style(run, size)
    run._element.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run = paragraph.add_run()
    set_citation_run_style(run, size)
    run._element.append(fld_sep)

    run = paragraph.add_run(number)
    set_citation_run_style(run, size)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    set_citation_run_style(run, size)
    run._element.append(fld_end)


def add_citation(paragraph, citation_text: str, size=Pt(9)):
    run = paragraph.add_run("[")
    set_citation_run_style(run, size)
    pos = 0
    for match in re.finditer(r"\d+", citation_text):
        if match.start() > pos:
            run = paragraph.add_run(citation_text[pos : match.start()])
            set_citation_run_style(run, size)
        add_ref_field(paragraph, match.group(0), size)
        pos = match.end()
    if pos < len(citation_text):
        run = paragraph.add_run(citation_text[pos:])
        set_citation_run_style(run, size)
    run = paragraph.add_run("]")
    set_citation_run_style(run, size)


def add_runs_with_reference_fields(paragraph, text: str, size=Pt(12)):
    pos = 0
    for match in CITATION_RE.finditer(text):
        if match.start() > pos:
            add_run_with_font(paragraph, text[pos : match.start()], size=size)
        add_citation(paragraph, match.group(1), size=Pt(9))
        pos = match.end()
    if pos < len(text):
        add_run_with_font(paragraph, text[pos:], size=size)


def set_cell_paragraph(cell, text: str):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    doc = cell.part.document
    set_paragraph_style(paragraph, doc, "表格文")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    add_runs_with_reference_fields(paragraph, text, size=Pt(10.5))


def cm_to_twips(cm: float) -> int:
    return int(round(cm * 567))


def set_table_col_widths(table, widths_cm):
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tbl = table._tbl
    tbl_pr = tbl.find(f"{{{ns}}}tblPr")
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(f"{{{ns}}}tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(cm_to_twips(sum(widths_cm))))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(f"{{{ns}}}tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.find(f"{{{ns}}}tblGrid")
    if tbl_grid is not None:
        for idx, col in enumerate(tbl_grid.findall(f"{{{ns}}}gridCol")):
            if idx < len(widths_cm):
                col.set(qn("w:w"), str(cm_to_twips(widths_cm[idx])))

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])


def insert_seq_field(paragraph, seq_name: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._element.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {seq_name} \\* ARABIC "
    paragraph.add_run()._element.append(instr)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._element.append(sep)

    number_run = paragraph.add_run("1")
    set_run_font(number_run)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._element.append(end)


def add_page_break(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


PAD = 0.46
SAFETY = 0.08


def text_display_width_cm(text: str) -> float:
    width = PAD
    for ch in text:
        width += 0.38 if ord(ch) > 127 else 0.25
    return width


def calc_col_widths(header, data_rows):
    cols = len(header)
    need = []
    short_cols = set()
    for col in range(cols):
        texts = [header[col]] + [row[col] for row in data_rows if col < len(row)]
        need.append(max(0.3, max(text_display_width_cm(text) for text in texts) + SAFETY))
        if max(len(text) for text in texts) <= 6:
            short_cols.add(col)

    total_need = sum(need)
    if total_need <= PAGE_W:
        return [round(value, 2) for value in need]

    widths = [0.0] * cols
    short_total = sum(need[col] for col in short_cols)
    long_cols = [col for col in range(cols) if col not in short_cols]
    available = PAGE_W - short_total
    if available <= 0 or not long_cols:
        ratio = PAGE_W / total_need
        widths = [value * ratio for value in need]
    else:
        for col in short_cols:
            widths[col] = need[col]
        long_total = sum(need[col] for col in long_cols)
        for col in long_cols:
            widths[col] = need[col] / long_total * available

    widths = [round(value, 2) for value in widths]
    if sum(widths) > PAGE_W + 0.01:
        ratio = PAGE_W / sum(widths)
        widths = [round(value * ratio, 2) for value in widths]
    return widths


def get_caption_before(lines, idx: int, prefix: str) -> str:
    pos = idx - 1
    while pos >= 0:
        text = lines[pos].strip()
        if not text:
            pos -= 1
            continue
        match = re.match(r"^" + prefix + r"\s*\d+[\-.]\d+\s*(.*)", text)
        return match.group(1).strip() if match else ""
    return ""


def resolve_asset(raw_path: str, asset_dir: Path) -> Path:
    cleaned = raw_path.strip().strip('"').strip("'")
    candidate = Path(cleaned)
    return candidate if candidate.is_absolute() else (asset_dir / candidate).resolve()


def add_image(doc, image_path: Path, caption: str, stats: dict, width_cm: float | None = None):
    target_width = min(width_cm or PAGE_W, PAGE_W)
    if image_path.exists() and image_path.is_file() and image_path.stat().st_size > 0:
        paragraph = doc.add_paragraph()
        set_paragraph_style(paragraph, doc, "图片")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(image_path), width=Cm(target_width))
        stats["images"] += 1
    else:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = add_run_with_font(paragraph, f"[缺失图片：{image_path.name}]", size=Pt(9))
        run.font.color.rgb = RGBColor(128, 128, 128)
        stats["missing_assets"].append(str(image_path))

    if caption:
        cap = doc.add_paragraph()
        set_paragraph_style(cap, doc, "Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_font(cap, "图", size=Pt(10.5))
        insert_seq_field(cap, "Figure")
        add_run_with_font(cap, f"-{caption}", size=Pt(10.5))
        stats["figure_captions"] += 1


def _frontmatter_end(lines: list[str]) -> int:
    if not lines or lines[0].strip() != "---":
        return 0
    for idx in range(1, len(lines)):
        if lines[idx].strip() == "---":
            return idx + 1
    return 0


def parse_and_write(
    md_path: Path,
    doc,
    *,
    asset_dir: Path,
    mermaid_images: list[Path],
    strip_heading_numbering: bool = True,
    chapter_page_breaks: bool = True,
):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    stats = {
        "headings": 0,
        "paragraphs": 0,
        "tables": 0,
        "images": 0,
        "figure_captions": 0,
        "table_captions": 0,
        "code_blocks": 0,
        "reference_entries": 0,
        "missing_assets": [],
    }

    i = _frontmatter_end(lines)
    first_heading = True
    in_references = False
    mermaid_index = 0

    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped == "---":
            i += 1
            continue

        if CAPTION_RE.match(stripped):
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            heading_text = heading_match.group(2).strip()
            if strip_heading_numbering:
                heading_text = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", heading_text)
            if level == 1 and chapter_page_breaks and not first_heading:
                add_page_break(doc)
            first_heading = False
            heading = doc.add_heading(heading_text, level=level)
            for run in heading.runs:
                set_run_font(run)
            in_references = bool(re.search(r"(参考文献|references)", heading_text, re.I))
            stats["headings"] += 1
            i += 1
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            add_image(doc, resolve_asset(image_match.group(2), asset_dir), image_match.group(1).strip(), stats)
            i += 1
            continue

        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1

            if lang == "mermaid":
                caption = ""
                caption_line = -1
                for pos in range(i, min(i + 5, len(lines))):
                    match = re.match(r"^图\s*\d+[\-.]\d+\s*(.*)", lines[pos].strip())
                    if match:
                        caption = match.group(1).strip()
                        caption_line = pos
                        break
                image_path = (
                    mermaid_images[mermaid_index]
                    if mermaid_index < len(mermaid_images)
                    else (asset_dir / f"mermaid_{mermaid_index + 1}.png")
                )
                mermaid_index += 1
                add_image(doc, image_path, caption, stats)
                if caption_line >= 0:
                    i = caption_line + 1
            else:
                paragraph = doc.add_paragraph()
                set_paragraph_style(paragraph, doc, "No Spacing")
                for idx, code_line in enumerate(code_lines):
                    run = paragraph.add_run(code_line + ("\n" if idx < len(code_lines) - 1 else ""))
                    set_run_font(run, cn_font="等线", en_font="Courier New", size=Pt(9))
                stats["code_blocks"] += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_start = i
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) < 2:
                continue

            header = [cell.strip() for cell in table_lines[0].split("|")[1:-1]]
            separator = bool(re.match(r"^\|?\s*:?-+", table_lines[1]))
            data_lines = table_lines[2:] if separator else table_lines[1:]
            data_rows = [[cell.strip() for cell in row.split("|")[1:-1]] for row in data_lines]
            if not header:
                continue

            caption = get_caption_before(lines, table_start, "表")
            cap = doc.add_paragraph()
            set_paragraph_style(cap, doc, "Caption")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_with_font(cap, "表", size=Pt(10.5))
            insert_seq_field(cap, "Table")
            if caption:
                add_run_with_font(cap, f"-{caption}", size=Pt(10.5))
            stats["table_captions"] += 1

            table = doc.add_table(rows=1 + len(data_rows), cols=len(header), style="Table Grid")
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            set_table_col_widths(table, calc_col_widths(header, data_rows))
            for col, value in enumerate(header):
                set_cell_paragraph(table.rows[0].cells[col], value)
            repeat_table_header(table.rows[0])
            for row_idx, row_data in enumerate(data_rows, 1):
                for col_idx, value in enumerate(row_data[: len(header)]):
                    set_cell_paragraph(table.rows[row_idx].cells[col_idx], value)
            stats["tables"] += 1
            continue

        ref_match = re.match(r"^\[(\d+)\]\s*(.*)", stripped)
        if ref_match:
            paragraph = doc.add_paragraph()
            add_bookmarked_reference_label(paragraph, ref_match.group(1))
            add_runs_with_reference_fields(paragraph, ref_match.group(2))
            stats["reference_entries"] += 1
            i += 1
            continue

        if re.match(r"^[-*]\s+(.*)", stripped):
            while i < len(lines):
                match = re.match(r"^[-*]\s+(.*)", lines[i].strip())
                if not match:
                    break
                paragraph = doc.add_paragraph()
                set_paragraph_style(paragraph, doc, "List Paragraph")
                add_run_with_font(paragraph, "• ")
                add_runs_with_reference_fields(paragraph, match.group(1))
                stats["paragraphs"] += 1
                i += 1
            continue

        if re.match(r"^(\d+)\.\s+(.*)", stripped):
            while i < len(lines):
                match = re.match(r"^(\d+)\.\s+(.*)", lines[i].strip())
                if not match:
                    break
                paragraph = doc.add_paragraph()
                if in_references:
                    add_bookmarked_reference_label(paragraph, match.group(1))
                    add_runs_with_reference_fields(paragraph, match.group(2))
                    stats["reference_entries"] += 1
                else:
                    set_paragraph_style(paragraph, doc, "List Paragraph")
                    add_runs_with_reference_fields(paragraph, f"（{match.group(1)}）{match.group(2)}")
                    stats["paragraphs"] += 1
                i += 1
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        add_runs_with_reference_fields(paragraph, stripped)
        stats["paragraphs"] += 1
        i += 1

    return stats


def validate_output(output_path: Path, expected_stats: dict, baseline: dict | None = None) -> dict:
    baseline = baseline or {"tables": 0, "images": 0}
    reopened = Document(str(output_path))
    expected_tables = baseline.get("tables", 0) + expected_stats["tables"]
    expected_images = baseline.get("images", 0) + expected_stats["images"]
    result = {
        "docx_reopen": True,
        "paragraph_count": len(reopened.paragraphs),
        "table_count": len(reopened.tables),
        "image_count": len(reopened.inline_shapes),
        "expected_table_count": expected_tables,
        "expected_image_count": expected_images,
        "checks": {
            "tables_preserved": len(reopened.tables) == expected_tables,
            "images_preserved": len(reopened.inline_shapes) == expected_images,
        },
    }
    result["pass"] = all(result["checks"].values())
    return result


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Convert Markdown to template-aware DOCX")
    parser.add_argument("-i", "--input", required=True, help="input Markdown file")
    parser.add_argument("-o", "--output", required=True, help="output .docx path")
    parser.add_argument("-t", "--template", help="optional .docx/.dotx template")
    parser.add_argument("--asset-dir", help="base directory for relative Markdown image paths; defaults to input directory")
    parser.add_argument("--mermaid-image", action="append", default=[], help="rendered Mermaid image in document order; repeat for multiple diagrams")
    parser.add_argument("--page-width-cm", type=float, help="usable page width; defaults to template/document margins")
    parser.add_argument("--keep-heading-numbering", action="store_true", help="do not strip manual numeric prefixes from headings")
    parser.add_argument("--no-chapter-page-breaks", action="store_true", help="do not insert page breaks before level-1 chapters")
    parser.add_argument("--strict-assets", action="store_true", help="fail when an image asset is missing")
    parser.add_argument("--report", help="optional JSON validation report path")
    return parser


def main() -> int:
    global PAGE_W
    args = build_parser().parse_args()
    input_path = Path(args.input).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()
    template_path = Path(args.template).expanduser().resolve() if args.template else None
    asset_dir = Path(args.asset_dir).expanduser().resolve() if args.asset_dir else input_path.parent
    mermaid_images = [Path(path).expanduser().resolve() for path in args.mermaid_image]

    if not input_path.is_file() or input_path.suffix.lower() not in {".md", ".markdown"}:
        raise SystemExit(f"invalid Markdown input: {input_path}")
    if output_path.suffix.lower() != ".docx":
        raise SystemExit("output must end with .docx")
    if template_path and (not template_path.is_file() or template_path.suffix.lower() not in {".docx", ".dotx"}):
        raise SystemExit(f"invalid template: {template_path}")
    if args.page_width_cm is not None and args.page_width_cm <= 0:
        raise SystemExit("--page-width-cm must be positive")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = load_document(template_path)
    PAGE_W = args.page_width_cm or usable_page_width_cm(doc)
    baseline = {"tables": len(doc.tables), "images": len(doc.inline_shapes)}

    stats = parse_and_write(
        input_path,
        doc,
        asset_dir=asset_dir,
        mermaid_images=mermaid_images,
        strip_heading_numbering=not args.keep_heading_numbering,
        chapter_page_breaks=not args.no_chapter_page_breaks,
    )

    if args.strict_assets and stats["missing_assets"]:
        raise SystemExit("missing assets: " + ", ".join(stats["missing_assets"]))

    doc.save(str(output_path))
    validation = validate_output(output_path, stats, baseline)
    report = {
        "input": str(input_path),
        "template": str(template_path) if template_path else None,
        "output": str(output_path),
        "page_width_cm": PAGE_W,
        "baseline": baseline,
        "stats": stats,
        "validation": validation,
    }

    if args.report:
        report_path = Path(args.report).expanduser().resolve()
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if validation["pass"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
